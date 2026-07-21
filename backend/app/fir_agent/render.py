"""
Renders a FIRDraft (schema.py) into a .docx file: a clearly-labeled DRAFT
document with the narrative, suggested sections (flagged unverified), an
evidence appendix, and the audit chain reference — so it is traceable back
to the exact AI decisions it was built from.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.fir_agent.schema import FIRDraft


def render_fir_draft_docx(draft: FIRDraft, output_path: str) -> str:
    doc = Document()

    # --- Watermark-style banner making draft status unmissable ---
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(f"DRAFT — {draft.status.value.upper()} — NOT A FILED DOCUMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    doc.add_heading("Cybercrime / Fraud Complaint — Draft", level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Draft ID: {draft.draft_id}\n").bold = True
    meta.add_run(f"Generated: {draft.created_at.isoformat()}\n")
    if draft.jurisdiction_police_station:
        meta.add_run(f"Jurisdiction (Police Station): {draft.jurisdiction_police_station}\n")
    if draft.jurisdiction_state:
        meta.add_run(f"State: {draft.jurisdiction_state}\n")
    meta.add_run(f"Audit chain reference: {draft.audit_chain_ref}\n")

    # --- Grounding warnings if any ---
    if draft.grounding_warnings:
        doc.add_heading("Grounding Warnings — Officer Must Review", level=2)
        warn_para = doc.add_paragraph()
        warn_run = warn_para.add_run(
            "The following claims in the AI-generated narrative could not be verified "
            "against the source evidence. They may be hallucinated and MUST be confirmed "
            "or removed before filing:\n\n"
        )
        warn_run.font.color.rgb = RGBColor(0xB0, 0x45, 0x00)
        for w in draft.grounding_warnings:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(w)
            r.bold = True
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    doc.add_heading("Complainant", level=2)
    c = draft.complainant
    doc.add_paragraph(f"Name: {c.full_name}")
    doc.add_paragraph(f"Contact: {c.contact_number}")
    if c.address:
        doc.add_paragraph(f"Address: {c.address}")

    doc.add_heading("Factual Narrative (AI-drafted — officer must review)", level=2)
    doc.add_paragraph(draft.narrative)

    doc.add_heading("Suggested Law Sections (UNVERIFIED — officer must confirm)", level=2)
    for s in draft.suggested_sections:
        p = doc.add_paragraph(style=None)
        p.add_run(f"{s.act} — {s.section}").bold = True
        p.add_run(f"\nBasis: {s.plain_language_basis}")
        p.add_run(f"\nVerified by officer: {'YES' if s.verified_by_officer else 'NO — PENDING REVIEW'}")

    doc.add_heading("Evidence Appendix", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
        "Event ID", "Timestamp", "Module", "Confidence", "Human Review Action"
    )
    for item in draft.evidence_items:
        row = table.add_row().cells
        row[0].text = item.event_id
        row[1].text = item.timestamp.isoformat()
        row[2].text = item.module_name.value
        row[3].text = f"{item.confidence_score:.2f}"
        row[4].text = item.review_action or "(not yet reviewed)"

    doc.add_heading("Disclaimer", level=2)
    disclaimer_p = doc.add_paragraph(draft.disclaimer)
    disclaimer_p.runs[0].italic = True

    doc.save(output_path)
    return output_path
